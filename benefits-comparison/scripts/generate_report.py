#!/usr/bin/env python3
"""Generate a Benefits Comparison Word document from structured JSON data.

The JSON schema is defined in the benefits-comparison SKILL.md (Phase 6).

Usage:
    python3 generate_report.py data.json
    python3 generate_report.py data.json --output "Smith_Benefits_2026-27.docx"

Requires:
    python3 -m pip install python-docx
"""
import json
import sys
import os
import argparse
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is not installed.")
    print("Run: python3 -m pip install python-docx")
    sys.exit(1)


# ── Color palette ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x1A, 0x73, 0x48)
RED   = RGBColor(0xC0, 0x39, 0x2B)
BLUE  = RGBColor(0x1A, 0x5C, 0x8A)
GRAY  = RGBColor(0xE8, 0xE8, 0xE8)
AMBER = RGBColor(0xD3, 0x6A, 0x00)


# ── Formatting helpers ────────────────────────────────────────────────────────

def shade_cell(cell, color: RGBColor) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color.red:02X}{color.green:02X}{color.blue:02X}")
    tc_pr.append(shd)


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(16)


def h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(13)


def h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(11)


def body(doc: Document, text: str, bold: bool = False, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color


def bullet(doc: Document, text: str, color: RGBColor = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def make_table(doc: Document, headers: list, rows: list, col_widths: list = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = str(h)
        shade_cell(cell, GRAY)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            if isinstance(val, tuple):
                text, color = val
                cell.text = str(text) if text is not None else "—"
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if color:
                            run.font.color.rgb = color
            else:
                cell.text = str(val) if val is not None else "—"
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            if i < len(table.columns):
                for row in table.rows:
                    row.cells[i].width = Inches(width)


def fmt_money(val) -> str:
    if val is None:
        return "—"
    try:
        f = float(val)
        return f"${f:,.0f}" if f == int(f) else f"${f:,.2f}"
    except (ValueError, TypeError):
        return str(val)


def fmt_tier_premium(tiers: dict, key: str) -> str:
    val = tiers.get(key)
    if val is None:
        return "—"
    unit = tiers.get("tier_unit", "biweekly")
    unit_abbr = {"biweekly": "biwk", "monthly": "mo", "semi-monthly": "semi-mo"}.get(unit, unit[:4])
    try:
        return f"${float(val):.2f}/{unit_abbr}"
    except (ValueError, TypeError):
        return str(val)


# ── Section builders ──────────────────────────────────────────────────────────

def add_family_section(doc: Document, fam: dict) -> None:
    members = fam.get("members", [])
    if not members:
        return
    h2(doc, "Family Members")
    rows = [
        (
            m.get("name", ""),
            m.get("relationship", ""),
            str(m.get("age", m.get("dob", ""))),
            m.get("notes", ""),
        )
        for m in members
    ]
    make_table(doc, ["Name", "Relationship", "Age / DOB", "Notes"], rows, [1.5, 2.0, 1.2, 2.8])


def add_employer_section(doc: Document, employer: dict) -> None:
    h2(doc, employer.get("name", "Employer"))
    body(
        doc,
        f"Plan Year: {employer.get('plan_year', '—')}  ·  "
        f"Pay Periods: {employer.get('pay_periods_per_year', 26)}/year  ·  "
        f"Employee: {employer.get('primary_employee', '—')}",
    )

    # Medical plans
    med_plans = employer.get("medical_plans", [])
    if med_plans:
        h3(doc, "Medical Plans")
        headers = ["Plan", "Deductible (Ind/Fam)", "OOP Max (Ind/Fam)", "PCP Copay", "EO", "EE+Spouse", "EE+Kids", "Family"]
        rows = []
        for p in med_plans:
            tiers = p.get("premiums_by_tier", {})
            rows.append([
                p.get("name", ""),
                f"{fmt_money(p.get('deductible_individual'))} / {fmt_money(p.get('deductible_family'))}",
                f"{fmt_money(p.get('oop_max_individual'))} / {fmt_money(p.get('oop_max_family'))}",
                p.get("pcp_copay", "—"),
                fmt_tier_premium(tiers, "employee_only"),
                fmt_tier_premium(tiers, "employee_spouse"),
                fmt_tier_premium(tiers, "employee_children"),
                fmt_tier_premium(tiers, "family"),
            ])
        make_table(doc, headers, rows)
        hsa_contrib = employer.get("employer_hsa_contribution_annual")
        if hsa_contrib:
            body(doc, f"Employer HSA contribution: {fmt_money(hsa_contrib)}/year (for HDHP-enrolled employees)", color=BLUE)

    # Dental
    dental = employer.get("dental")
    if dental:
        h3(doc, f"Dental — {dental.get('carrier', '')}")
        rows = [
            ("Annual maximum", fmt_money(dental.get("annual_max_per_person")) + "/person"),
            ("Deductible", f"${dental.get('deductible_individual', 50)} ind / ${dental.get('deductible_family', 150)} fam"),
            ("Preventive (Class I)", f"{dental.get('preventive_pct', 100)}%"),
            ("Basic / Restorative (Class II)", f"{dental.get('basic_pct', 0)}%"),
            ("Major — crowns (Class III)", f"{dental.get('major_pct', 0)}%"),
            ("Orthodontics", dental.get("ortho_benefit", "None listed")),
            ("Waiting periods", dental.get("waiting_periods", "None listed")),
        ]
        make_table(doc, ["Benefit", "Detail"], rows, [2.5, 4.0])
        # Dental premiums
        d_tiers = dental.get("premiums_by_tier", {})
        if d_tiers:
            tier_rows = [(k.replace("_", " ").title(), fmt_tier_premium(d_tiers, k)) for k in d_tiers if k != "tier_unit"]
            if tier_rows:
                doc.add_paragraph()
                make_table(doc, ["Tier", "Premium"], tier_rows, [3.0, 3.5])

    # Vision
    vision = employer.get("vision")
    if vision:
        h3(doc, f"Vision — {vision.get('carrier', '')}")
        rows = [
            ("Exam copay", f"${vision.get('exam_copay', 10)} / 12 months"),
            ("Frame allowance (standard / featured)", f"${vision.get('frame_allowance_standard', 0)} / ${vision.get('frame_allowance_featured', 0)}"),
            ("Contact lens allowance", f"${vision.get('contacts_allowance', 0)} + up to $60 fitting"),
            ("Standard progressive", "$0"),
            ("Premium / custom progressive", f"${vision.get('premium_progressive_copay', '—')} copay"),
        ]
        make_table(doc, ["Benefit", "Detail"], rows, [2.5, 4.0])
        v_tiers = vision.get("premiums_by_tier", {})
        if v_tiers:
            tier_rows = [(k.replace("_", " ").title(), fmt_tier_premium(v_tiers, k)) for k in v_tiers if k != "tier_unit"]
            if tier_rows:
                doc.add_paragraph()
                make_table(doc, ["Tier", "Premium"], tier_rows, [3.0, 3.5])


def add_scenario_section(doc: Document, scenario: dict) -> None:
    is_rec = scenario.get("recommended", False)
    name = scenario.get("name", "Scenario")

    if is_rec:
        h2(doc, f"★  {name}  — RECOMMENDED")
    else:
        h2(doc, name)

    desc = scenario.get("description", "")
    if desc:
        body(doc, desc)

    # Assignments
    assignments = scenario.get("assignments", {})
    if assignments:
        h3(doc, "Enrollment Assignments")
        make_table(doc, ["Family Member", "Plan"], list(assignments.items()), [2.0, 4.5])

    # Cost breakdown
    h3(doc, "Annual Cost Breakdown")
    cost_rows = [("Annual Premiums", fmt_money(scenario.get("annual_premiums")))]

    hsa_employer = scenario.get("hsa_employer_contribution", 0) or 0
    if hsa_employer:
        cost_rows.append(("Employer HSA Contribution", f"−{fmt_money(hsa_employer)}"))

    hsa_savings = scenario.get("hsa_tax_savings", 0) or 0
    if hsa_savings:
        cost_rows.append(("HSA Tax Savings (est.)", f"−{fmt_money(hsa_savings)}"))

    dcfsa_savings = scenario.get("dcfsa_tax_savings", 0) or 0
    if dcfsa_savings:
        cost_rows.append(("Dep. Care FSA Tax Savings (est.)", f"−{fmt_money(dcfsa_savings)}"))

    cost_rows.append(("Estimated Out-of-Pocket", fmt_money(scenario.get("estimated_oop"))))

    net = scenario.get("net_annual_cost")
    net_color = GREEN if is_rec else None
    cost_rows.append(("Net Annual Cost", (fmt_money(net), net_color)))
    make_table(doc, ["Item", "Amount"], cost_rows, [4.0, 2.5])

    # Pros / Cons
    pros = scenario.get("pros", [])
    cons = scenario.get("cons", [])
    if pros or cons:
        h3(doc, "Pros & Cons")
        for pro in pros:
            bullet(doc, f"✓  {pro}", GREEN)
        for con in cons:
            bullet(doc, f"✗  {con}", RED)

    note = scenario.get("recommendation_note", "")
    if note:
        doc.add_paragraph()
        body(doc, note, bold=True, color=GREEN if is_rec else BLUE)


# ── Main report builder ───────────────────────────────────────────────────────

def generate_report(data: dict, output_path: str) -> None:
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title_p = doc.add_heading(data.get("report_title", "Benefits Comparison"), level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(20)

    fam = data.get("family", {})
    subtitle_parts = [fam.get("location", ""), f"Generated {data.get('generated_date', date.today().isoformat())}"]
    sub = doc.add_paragraph("  ·  ".join(p for p in subtitle_parts if p))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_family_section(doc, fam)
    doc.add_page_break()

    # ── Rate Reference Tables ─────────────────────────────────────────────────
    h1(doc, "Plan Rate Reference")
    for employer in data.get("employers", []):
        add_employer_section(doc, employer)
        doc.add_paragraph()
    doc.add_page_break()

    # ── Scenario Rundowns ─────────────────────────────────────────────────────
    h1(doc, "Enrollment Scenarios")
    for scenario in data.get("scenarios", []):
        add_scenario_section(doc, scenario)
        doc.add_paragraph()
    doc.add_page_break()

    # ── Master Comparison Table ───────────────────────────────────────────────
    h1(doc, "Master Comparison")
    scenarios = data.get("scenarios", [])
    if scenarios:
        headers = ["Scenario", "Annual Premium", "Est. OOP", "Tax Savings", "Net Annual Cost", "Rec?"]
        rows = []
        for s in scenarios:
            tax_sav = (s.get("hsa_tax_savings") or 0) + (s.get("dcfsa_tax_savings") or 0) + (s.get("hsa_employer_contribution") or 0)
            is_rec = s.get("recommended", False)
            rows.append([
                s.get("name", ""),
                fmt_money(s.get("annual_premiums")),
                fmt_money(s.get("estimated_oop")),
                fmt_money(tax_sav) if tax_sav else "—",
                (fmt_money(s.get("net_annual_cost")), GREEN if is_rec else None),
                ("★ Yes", GREEN) if is_rec else "No",
            ])
        make_table(doc, headers, rows, [2.2, 1.2, 1.0, 1.1, 1.3, 0.7])
    doc.add_page_break()

    # ── Dental / Vision Comparison ────────────────────────────────────────────
    dv_notes = data.get("dental_vision_comparison")
    if dv_notes:
        h1(doc, "Dental & Vision Comparison")
        body(doc, dv_notes)
        doc.add_page_break()

    # ── Recommendation & Action Items ────────────────────────────────────────
    h1(doc, "Recommendation & Action Items")
    rec = data.get("recommendation", {})
    if rec:
        body(doc, rec.get("headline", ""), bold=True, color=GREEN)
        doc.add_paragraph()
        body(doc, rec.get("reasoning", ""))
        doc.add_paragraph()

        action_items = rec.get("action_items", [])
        if action_items:
            h2(doc, "Action Items")
            rows = [
                (a.get("owner", ""), a.get("task", ""), a.get("deadline", ""))
                for a in action_items
            ]
            make_table(doc, ["Owner", "Task", "Deadline"], rows, [1.2, 4.0, 1.5])

    # ── Additional Notes ──────────────────────────────────────────────────────
    extra = data.get("additional_notes")
    if extra:
        doc.add_paragraph()
        h2(doc, "Additional Notes")
        body(doc, extra)

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a Benefits Comparison Word document from JSON data.")
    parser.add_argument("data_json", help="Path to the data.json file (see SKILL.md for schema)")
    parser.add_argument(
        "--output",
        default="benefits_comparison.docx",
        help="Output .docx file path (default: benefits_comparison.docx)",
    )
    args = parser.parse_args()

    if not os.path.exists(args.data_json):
        print(f"Error: data file not found: {args.data_json}")
        sys.exit(1)

    with open(args.data_json, encoding="utf-8") as f:
        data = json.load(f)

    generate_report(data, args.output)
